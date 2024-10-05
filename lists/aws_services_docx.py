from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('AWS Services for Web Development, CI/CD, Data Analytics, and Deployment', 0)

# Web Development section
doc.add_heading('1. Web Development', level=1)
web_services = [
    "1. Amazon EC2 (Elastic Compute Cloud)",
    "2. Amazon S3 (Simple Storage Service)",
    "3. Amazon RDS (Relational Database Service)",
    "4. Amazon DynamoDB",
    "5. Amazon Route 53",
    "6. AWS Lambda",
    "7. Amazon API Gateway",
    "8. Amazon Lightsail",
    "9. AWS Elastic Beanstalk",
    "10. AWS Amplify"
]

for service in web_services:
    doc.add_paragraph(service)

# CI/CD section
doc.add_heading('2. CI/CD (Continuous Integration/Continuous Deployment)', level=1)
cicd_services = [
    "11. AWS CodePipeline",
    "12. AWS CodeBuild",
    "13. AWS CodeDeploy",
    "14. AWS CodeCommit",
    "15. AWS CodeArtifact",
    "16. AWS CloudFormation",
    "17. AWS CodeStar",
    "18. AWS OpsWorks",
    "19. AWS Cloud9",
    "20. AWS Systems Manager"
]

for service in cicd_services:
    doc.add_paragraph(service)

# Data Analytics section
doc.add_heading('3. Data Analytics', level=1)
data_services = [
    "21. Amazon Redshift",
    "22. Amazon Athena",
    "23. Amazon Kinesis",
    "24. AWS Glue",
    "25. Amazon QuickSight",
    "26. AWS Data Pipeline",
    "27. Amazon EMR (Elastic MapReduce)",
    "28. AWS Lake Formation",
    "29. Amazon Elasticsearch Service",
    "30. AWS Sagemaker (For machine learning and AI)"
]

for service in data_services:
    doc.add_paragraph(service)

# Deployment and Management section
doc.add_heading('4. Deployment and Management', level=1)
deployment_services = [
    "31. Amazon ECR (Elastic Container Registry)",
    "32. Amazon ECS (Elastic Container Service)",
    "33. Amazon EKS (Elastic Kubernetes Service)",
    "34. AWS Fargate",
    "35. AWS CloudWatch",
    "36. AWS IAM (Identity and Access Management)",
    "37. AWS Step Functions",
    "38. Amazon CloudFront",
    "39. AWS Elastic Load Balancing (ELB)",
    "40. Amazon Inspector"
]

for service in deployment_services:
    doc.add_paragraph(service)

# Save the document as a Word file
file_path = "C:\\Users\\HP\\Desktop\\studymaterial\\scripts_coding\\resourses_aws.docx"

doc.save(file_path)

file_path
